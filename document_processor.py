from typing import Union
from langchain_core.documents import Document
from fastapi import UploadFile
from datetime import datetime
import pdfplumber
from docx import Document as DocxDocument
import io

async def process_uploaded_file(file: UploadFile) -> Union[Document, None]:
    filename = file.filename.lower()
    content = await file.read()
    
    try:
        if filename.endswith('.txt'):
            text = content.decode('utf-8')
        elif filename.endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join(page.extract_text() for page in pdf.pages)
        elif filename.endswith('.docx'):
            doc = DocxDocument(io.BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs)
        else:
            return None
            
        return Document(
            page_content=text,
            metadata={
                "source": file.filename,
                "type": "uploaded_file",
                "upload_time": str(datetime.now())
            }
        )
    except Exception as e:
        print(f"Error processing file: {e}")
        return None